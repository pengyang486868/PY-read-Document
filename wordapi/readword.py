from docx import Document
import utils
import os
import re


def readtxt(full_path):
    doc = Document(full_path)
    result = []
    for para in doc.paragraphs:
        if len(para.text) < 1:
            continue
        result.append(para.text)

    # table text
    maxrows = 500
    rowcount = 0
    for tb in doc.tables:
        try:
            rowcount += len(tb.rows)
            if rowcount > maxrows:
                break
            for row in tb.rows:
                for cell in row.cells:
                    if len(cell.text) > 0:
                        result.append(cell.text)
        except:
            continue
    return result


def readimg(full_path, savedir, save_prefix=''):
    if not savedir:
        return
    savedir = os.path.join(savedir, utils.generate_dirname())
    os.mkdir(savedir)

    doc = Document(full_path)
    doc_part = doc.part
    imgcount = 0

    # delete empty paragraphs
    notemptyp = []
    for p in doc.paragraphs:
        xmlstr = p._p.xml
        if 'graphicData' in xmlstr or len(p.text.split()) > 0:
            notemptyp.append(p)

    contextdic = {}
    last_text = ''

    # table text
    maxrows = 500
    rowcount = 0
    for tb in doc.tables:
        try:
            rowcount += len(tb.rows)
            if rowcount > maxrows:
                break
            for row in tb.rows:
                for cell in row.cells:
                    if len(cell.text) == 0:
                        continue
                    for indx, p in enumerate(cell.paragraphs):
                        xmlstr = p._p.xml
                        if 'graphicData' not in xmlstr:
                            if len(p.text.split()) > 0:
                                last_text = p.text
                            continue

                        ridreg = re.compile(r'r:embed="(.*)"')
                        cur_rids = ridreg.findall(xmlstr)
                        if len(cur_rids) < 1:
                            continue
                        contextdic[cur_rids[0]] = last_text
        except Exception as e:
            continue

    # give text to graphic paragraphs
    for indx, p in enumerate(notemptyp):
        xmlstr = p._p.xml
        if 'graphicData' not in xmlstr:
            if len(p.text.split()) > 0:
                last_text = p.text

        # format: ===r:embed="rId5"===
        ridreg = re.compile(r'r:embed="(.*)"')
        cur_rids = ridreg.findall(xmlstr)
        if len(cur_rids) < 1:
            continue
        if notemptyp[indx + 1].text.startswith('图'):
            contextdic[cur_rids[0]] = notemptyp[indx + 1].text
        elif indx < 1:
            contextdic[cur_rids[0]] = os.path.split(full_path)[1]
        else:
            contextdic[cur_rids[0]] = last_text

    # fill image info
    imginfo = []
    for ishape in doc.inline_shapes:
        imgcount += 1
        blip = ishape._inline.graphic.graphicData.pic.blipFill.blip
        rID = blip.embed
        image_part = doc_part.related_parts[rID]

        # save
        fname = save_prefix + '-' + str(imgcount) + '-' + rID + '.png'
        fr = open(os.path.join(savedir, fname), "wb")
        fr.write(image_part.blob)
        fr.close()

        # find related text
        imginfo.append({'fname': os.path.join(savedir, fname),
                        'relatedtxt': contextdic[rID] if rID in contextdic else ''})

    return imginfo
